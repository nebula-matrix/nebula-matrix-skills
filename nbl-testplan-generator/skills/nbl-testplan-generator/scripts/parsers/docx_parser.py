#!/usr/bin/env python3
"""docx_parser.py - Parse functional specification .docx into structured JSON."""
import json
import re
import sys
from pathlib import Path
from docx import Document


def parse_docx(filepath):
    """Parse a .docx functional specification into a hierarchical dictionary.

    Returns a dict with keys:
        document_title, total_paragraphs, total_tables,
        feature_ids, sections, tables
    """
    doc = Document(filepath)
    filepath = Path(filepath)
    feature_pattern = re.compile(r"【(\w+\.\d+)】")

    # First pass: collect all paragraph metadata
    para_data = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name
        features = feature_pattern.findall(text)
        para_data.append({
            "index": i,
            "style": style,
            "text": text,
            "feature_ids": features,
        })

    # Extract all tables
    all_tables = []
    for idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        all_tables.append({"table_index": idx, "rows": rows})

    result = {
        "document_title": filepath.stem,
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "feature_ids": [],
        "sections": [],
        "tables": all_tables,
    }

    # Build hierarchical section structure
    current_section = None
    current_subsection = None
    current_subsubsection = None
    current_content = []
    current_sub_content = []
    current_subsub_content = []

    def flush_content():
        nonlocal current_content, current_sub_content, current_subsub_content
        if current_subsubsection is not None:
            current_subsubsection["content"] = current_subsub_content
            current_subsub_content = []
        elif current_subsection is not None:
            current_subsection["content"] = current_sub_content
            current_sub_content = []
        elif current_section is not None:
            current_section["content"] = current_content
            current_content = []

    for pd_item in para_data:
        style = pd_item["style"]
        text = pd_item["text"]
        features = pd_item["feature_ids"]
        if not text and not features:
            continue

        if style == "Heading 1":
            flush_content()
            current_section = {
                "level": 1,
                "title": text,
                "children": [],
                "content": [],
            }
            result["sections"].append(current_section)
            current_subsection = None
            current_subsubsection = None
        elif style == "Heading 2":
            if current_section is None:
                continue
            if current_subsubsection is not None and current_subsection is not None:
                current_subsection["children"].append(current_subsubsection)
                current_subsubsection = None
            if current_subsection is not None:
                current_subsection["content"] = current_sub_content
                current_sub_content = []
                current_section["children"].append(current_subsection)
            current_subsection = {
                "level": 2,
                "title": text,
                "children": [],
                "content": [],
            }
            current_subsubsection = None
        elif style == "Heading 3":
            if current_subsection is None:
                continue
            if current_subsubsection is not None:
                current_subsection["children"].append(current_subsubsection)
            current_subsubsection = {"level": 3, "title": text, "content": []}
        else:
            entry = {"text": text}
            if features:
                entry["feature_ids"] = features
                result["feature_ids"].extend(features)
            if style != "Normal":
                entry["style"] = style
            if current_subsubsection is not None:
                current_subsub_content.append(entry)
            elif current_subsection is not None:
                current_sub_content.append(entry)
            elif current_section is not None:
                current_content.append(entry)

    # Flush remaining content at end of document
    if current_subsubsection is not None and current_subsection is not None:
        current_subsubsection["content"] = current_subsub_content
        current_subsection["children"].append(current_subsubsection)
    if current_subsection is not None and current_section is not None:
        current_subsection["content"] = current_sub_content
        current_section["children"].append(current_subsection)
    if current_section is not None:
        current_section["content"] = current_content

    return result


def get_section_by_title(parsed, title_keyword):
    """Find a section (at any depth) whose title contains the keyword."""
    for section in parsed["sections"]:
        if title_keyword in section["title"]:
            return section
        for child in section.get("children", []):
            if title_keyword in child["title"]:
                return child
            for grandchild in child.get("children", []):
                if title_keyword in grandchild["title"]:
                    return grandchild
    return None


def get_features_section(parsed):
    """Return the '模块特性' section."""
    return get_section_by_title(parsed, "模块特性")


def get_functional_details_section(parsed):
    """Return the '功能详述' section."""
    return get_section_by_title(parsed, "功能详述")


def get_parameters_section(parsed):
    """Return the '参数列表' section."""
    return get_section_by_title(parsed, "参数列表")


def extract_feature_map(features_section):
    """Extract a mapping of feature_id -> {description, raw_text} from a section."""
    feature_map = {}
    if not features_section:
        return feature_map

    def _extract(content_list):
        for entry in content_list:
            for fid in entry.get("feature_ids", []):
                desc = re.sub(r"【\w+\.\d+】", "", entry["text"]).strip()
                feature_map[fid] = {
                    "description": desc,
                    "raw_text": entry["text"],
                }

    _extract(features_section.get("content", []))
    for child in features_section.get("children", []):
        _extract(child.get("content", []))
        for grandchild in child.get("children", []):
            _extract(grandchild.get("content", []))
    return feature_map


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python docx_parser.py <input.docx> [output.json]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    result = parse_docx(input_path)

    print(f"Document: {result['document_title']}")
    print(f"Paragraphs: {result['total_paragraphs']}")
    print(f"Tables: {result['total_tables']}")
    print(f"Sections: {[s['title'] for s in result['sections']]}")
    print(f"Feature IDs: {result['feature_ids']}")

    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"Output written to: {output_path}")
